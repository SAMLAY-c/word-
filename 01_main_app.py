import tkinter as tk
from tkinter import filedialog, messagebox, ttk, colorchooser, simpledialog, scrolledtext
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm # Added Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
import os
import threading
import json
import requests
import re
import configparser
import time 

CONFIG_FILE = "config.ini" # Define config file name as a constant

class DocxFormatter:
    def __init__(self, root):
        self.root = root
        self.root.title("Word文档格式化工具") 
        self.root.geometry("1000x780") # Increased height slightly for new button and default config load status
        
        # Style variables - Initialize with hardcoded defaults first
        self.title_font = tk.StringVar(value="黑体") 
        self.title_size = tk.IntVar(value=22)
        self.title_color_val = tk.StringVar(value="#000000") 
        self.title_bold = tk.BooleanVar(value=True)
        
        self.h1_font = tk.StringVar(value="黑体") 
        self.h1_size = tk.IntVar(value=18)
        self.h1_color_val = tk.StringVar(value="#000000")
        self.h1_bold = tk.BooleanVar(value=True)
        
        self.h2_font = tk.StringVar(value="楷体") 
        self.h2_size = tk.IntVar(value=16)
        self.h2_color_val = tk.StringVar(value="#000000")
        self.h2_bold = tk.BooleanVar(value=True)
        
        self.h3_font = tk.StringVar(value="宋体") 
        self.h3_size = tk.IntVar(value=14)
        self.h3_color_val = tk.StringVar(value="#000000")
        self.h3_bold = tk.BooleanVar(value=True)
        
        self.normal_font = tk.StringVar(value="仿宋") 
        self.normal_size = tk.IntVar(value=12)
        self.normal_color_val = tk.StringVar(value="#000000")
        self.normal_bold = tk.BooleanVar(value=False)
        self.indent_chars_str = tk.StringVar(value="2") # For indent entry, will be get/set as string

        self.toc_title = tk.StringVar(value="目 录") 
        self.filename = tk.StringVar()
        
        self.document_sections = []
        self.document_title = tk.StringVar(value="公文标题示例") 
        
        self.deepseek_api_key = tk.StringVar()
        self.deepseek_model = tk.StringVar(value="deepseek-chat") 
        
        self.api_dependencies_status = tk.StringVar(value="未检查") 
        
        self.log_text = None 
        self.progress_bar = None
        self.color_previews = {}
        self.analyze_button = None 

        self.api_session = requests.Session()

        self.add_logo = tk.BooleanVar(value=False)
        self.logo_path = tk.StringVar(value="")
        self.logo_position = tk.StringVar(value="left") 
        self.logo_width_cm = tk.DoubleVar(value=2.5) 

        # Load API settings first, then UI settings which might depend on config file structure
        self.load_api_settings() 
        self.load_default_ui_settings() # Load UI defaults after variables are initialized

        self.create_widgets() # Create widgets after settings are potentially loaded
        self.check_and_install_dependencies()
        
    def load_api_settings(self):
        """从配置文件加载 API Key 和模型选择"""
        config = configparser.ConfigParser()
        if os.path.exists(CONFIG_FILE):
            config.read(CONFIG_FILE)
            if "DEEPSEEK" in config:
                self.deepseek_api_key.set(config["DEEPSEEK"].get("api_key", ""))
                self.deepseek_model.set(config["DEEPSEEK"].get("model", "deepseek-chat"))
        else: 
            # If config file doesn't exist, it will be created by save_api_settings or save_default_ui_settings
            self.deepseek_api_key.set("") # Ensure variables have initial values
            self.deepseek_model.set("deepseek-chat")


    def save_api_settings(self):
        """保存 API Key 和模型选择到配置文件"""
        config = configparser.ConfigParser()
        if os.path.exists(CONFIG_FILE):
            config.read(CONFIG_FILE) 
        
        if "DEEPSEEK" not in config:
            config["DEEPSEEK"] = {}
        config["DEEPSEEK"]["api_key"] = self.deepseek_api_key.get()
        config["DEEPSEEK"]["model"] = self.deepseek_model.get()
        
        try:
            with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
                config.write(f)
            messagebox.showinfo("成功", "API 设置已保存到 config.ini")
        except Exception as e:
            messagebox.showerror("错误", f"保存API设置失败: {e}")
            self.log(f"保存API设置失败: {e}")

    # --- New methods for loading and saving default UI settings ---
    def load_default_ui_settings(self):
        """从配置文件加载默认UI设置"""
        config = configparser.ConfigParser()
        if not os.path.exists(CONFIG_FILE):
            self.log("配置文件 config.ini 未找到，使用程序内置默认UI设置。")
            return

        try:
            config.read(CONFIG_FILE, encoding='utf-8')
        except Exception as e:
            self.log(f"读取配置文件 config.ini 失败: {e}。使用程序内置默认UI设置。")
            return

        if "DEFAULT_UI_SETTINGS" not in config:
            self.log("未找到默认UI配置区域，使用程序内置默认UI设置。")
            return

        settings = config["DEFAULT_UI_SETTINGS"]
        
        def get_setting(key, var, default_val, var_type="string"):
            try:
                if var_type == "string": var.set(settings.get(key, default_val))
                elif var_type == "int": var.set(settings.getint(key, fallback=default_val))
                elif var_type == "bool": var.set(settings.getboolean(key, fallback=default_val))
                elif var_type == "float": var.set(settings.getfloat(key, fallback=default_val)) # For DoubleVar
            except (configparser.NoOptionError, ValueError) as e:
                var.set(default_val)
                self.log(f"加载默认设置 '{key}' 失败或值无效 ({e})，使用内置值: {default_val}")
            except Exception as e: # Catch any other unexpected errors during get
                var.set(default_val)
                self.log(f"加载默认设置 '{key}' 时发生未知错误 ({e})，使用内置值: {default_val}")


        get_setting("title_font", self.title_font, "黑体")
        get_setting("title_size", self.title_size, 22, "int")
        get_setting("title_color", self.title_color_val, "#000000")
        get_setting("title_bold", self.title_bold, True, "bool")

        get_setting("h1_font", self.h1_font, "黑体")
        get_setting("h1_size", self.h1_size, 18, "int")
        get_setting("h1_color", self.h1_color_val, "#000000")
        get_setting("h1_bold", self.h1_bold, True, "bool")

        get_setting("h2_font", self.h2_font, "楷体")
        get_setting("h2_size", self.h2_size, 16, "int")
        get_setting("h2_color", self.h2_color_val, "#000000")
        get_setting("h2_bold", self.h2_bold, True, "bool")

        get_setting("h3_font", self.h3_font, "宋体")
        get_setting("h3_size", self.h3_size, 14, "int")
        get_setting("h3_color", self.h3_color_val, "#000000")
        get_setting("h3_bold", self.h3_bold, True, "bool")

        get_setting("normal_font", self.normal_font, "仿宋")
        get_setting("normal_size", self.normal_size, 12, "int")
        get_setting("normal_color", self.normal_color_val, "#000000")
        get_setting("normal_bold", self.normal_bold, False, "bool")
        
        get_setting("indent_chars", self.indent_chars_str, "2") # Loaded as string for Entry

        get_setting("toc_title", self.toc_title, "目 录")
        # self.document_title is usually document-specific, so not saving/loading as default.

        get_setting("add_logo", self.add_logo, False, "bool")
        # Do not load logo_path as default, user should select it each time if needed, or it might be invalid.
        # self.logo_path.set(settings.get("logo_path", "")) # Or clear it: self.logo_path.set("")
        get_setting("logo_position", self.logo_position, "left")
        get_setting("logo_width_cm", self.logo_width_cm, 2.5, "float")
        
        self.log("已加载默认UI配置。")
        # Update color previews after loading settings
        if hasattr(self, 'color_previews') and self.color_previews: # Check if widgets are created
            for key, var_name in [("title_color", "title_color_val"), ("h1_color", "h1_color_val"), 
                                  ("h2_color", "h2_color_val"), ("h3_color", "h3_color_val"), 
                                  ("normal_color", "normal_color_val")]:
                if key in self.color_previews:
                    color_val = getattr(self, var_name).get()
                    self.color_previews[key].config(bg=color_val)


    def save_default_ui_settings(self):
        """保存当前UI设置为默认配置"""
        config = configparser.ConfigParser()
        if os.path.exists(CONFIG_FILE):
            try:
                config.read(CONFIG_FILE, encoding='utf-8')
            except Exception as e:
                messagebox.showerror("错误", f"读取现有配置文件失败: {e}")
                self.log(f"读取现有配置文件失败: {e}")
                return

        if "DEFAULT_UI_SETTINGS" not in config:
            config["DEFAULT_UI_SETTINGS"] = {}
        
        settings = config["DEFAULT_UI_SETTINGS"]
        settings["title_font"] = self.title_font.get()
        settings["title_size"] = str(self.title_size.get())
        settings["title_color"] = self.title_color_val.get()
        settings["title_bold"] = str(self.title_bold.get())

        settings["h1_font"] = self.h1_font.get()
        settings["h1_size"] = str(self.h1_size.get())
        settings["h1_color"] = self.h1_color_val.get()
        settings["h1_bold"] = str(self.h1_bold.get())

        settings["h2_font"] = self.h2_font.get()
        settings["h2_size"] = str(self.h2_size.get())
        settings["h2_color"] = self.h2_color_val.get()
        settings["h2_bold"] = str(self.h2_bold.get())

        settings["h3_font"] = self.h3_font.get()
        settings["h3_size"] = str(self.h3_size.get())
        settings["h3_color"] = self.h3_color_val.get()
        settings["h3_bold"] = str(self.h3_bold.get())

        settings["normal_font"] = self.normal_font.get()
        settings["normal_size"] = str(self.normal_size.get())
        settings["normal_color"] = self.normal_color_val.get()
        settings["normal_bold"] = str(self.normal_bold.get())
        
        settings["indent_chars"] = self.indent_entry.get() # Save from indent_entry directly

        settings["toc_title"] = self.toc_title.get()

        settings["add_logo"] = str(self.add_logo.get())
        # Do not save logo_path as default.
        # settings["logo_path"] = self.logo_path.get() 
        settings["logo_position"] = self.logo_position.get()
        settings["logo_width_cm"] = str(self.logo_width_cm.get())

        try:
            with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
                config.write(f)
            messagebox.showinfo("成功", "当前UI配置已保存为默认设置。\n下次启动时将自动加载。")
            self.log("默认UI配置已保存。")
        except Exception as e:
            messagebox.showerror("错误", f"保存默认UI配置失败: {e}")
            self.log(f"保存默认UI配置失败: {e}")
    # --- End of new methods ---
            
    def check_and_install_dependencies(self):
        """检查依赖项，并在缺失时提示用户手动安装"""
        dependencies = [("docx", "python-docx"), ("requests", "requests")]
        missing_dependencies = []
        for import_name, install_name in dependencies:
            try:
                __import__(import_name)
                self.log(f"{install_name} 已安装。")
            except ImportError:
                self.log(f"{install_name} 未安装。请在激活虚拟环境后运行: pip install {install_name}")
                missing_dependencies.append(install_name)
            except Exception as e:
                self.log(f"检查 {install_name} 时发生未知错误: {str(e)}")
                missing_dependencies.append(f"{install_name} (检查失败)")

        if not missing_dependencies:
            self.api_dependencies_status.set("所有依赖已就绪")
        else:
            self.api_dependencies_status.set(f"缺失依赖: {', '.join(missing_dependencies)}")
            # Delay messagebox to ensure main window is fully up
            self.root.after(100, lambda: messagebox.showwarning("依赖缺失", 
                                   f"以下依赖包未能加载或缺失，请在激活虚拟环境后，通过 pip 手动安装:\n\n{', '.join(missing_dependencies)}\n\n例如: pip install python-docx requests"))

    def create_widgets(self):
        notebook = ttk.Notebook(self.root)
        notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        self.notebook = notebook
        
        basic_frame = ttk.Frame(notebook); notebook.add(basic_frame, text="基本设置")
        content_frame = ttk.Frame(notebook); notebook.add(content_frame, text="文档内容")
        
        header_footer_frame = ttk.Frame(notebook)
        notebook.add(header_footer_frame, text="页眉与页脚")
        self.setup_header_footer_tab(header_footer_frame)

        ai_frame = ttk.Frame(notebook); notebook.add(ai_frame, text="AI 标题识别")
        title_settings_frame = ttk.Frame(notebook); notebook.add(title_settings_frame, text="标题设置")
        styles_frame = ttk.Frame(notebook); notebook.add(styles_frame, text="字体设置")
        log_frame = ttk.Frame(notebook); notebook.add(log_frame, text="执行日志")
        
        # Basic Settings Tab
        basic_settings_content_frame = ttk.Frame(basic_frame)
        basic_settings_content_frame.pack(padx=10, pady=10, fill="x", anchor="n")

        ttk.Label(basic_settings_content_frame, text="文件名称(不需要输入.docx):", font=("Arial", 10)).grid(row=0, column=0, sticky="w", padx=10, pady=5) 
        ttk.Entry(basic_settings_content_frame, textvariable=self.filename, width=40).grid(row=0, column=1, sticky="ew", padx=10, pady=5)
        
        ttk.Label(basic_settings_content_frame, text="目录标题:", font=("Arial", 10)).grid(row=1, column=0, sticky="w", padx=10, pady=5) 
        ttk.Entry(basic_settings_content_frame, textvariable=self.toc_title, width=40).grid(row=1, column=1, sticky="ew", padx=10, pady=5)
        
        ttk.Label(basic_settings_content_frame, text="文档标题:", font=("Arial", 10)).grid(row=2, column=0, sticky="w", padx=10, pady=5) 
        ttk.Entry(basic_settings_content_frame, textvariable=self.document_title, width=40).grid(row=2, column=1, sticky="ew", padx=10, pady=5)
        basic_settings_content_frame.columnconfigure(1, weight=1) # Make entry expand

        # --- Add Save as Default Button ---
        save_default_btn = ttk.Button(basic_settings_content_frame, text="保存当前配置为默认", command=self.save_default_ui_settings)
        save_default_btn.grid(row=0, column=2, rowspan=3, padx=20, pady=5, sticky="ns")
        # --- End Add Save as Default Button ---

        api_frame = ttk.LabelFrame(basic_frame, text="DeepSeek API 设置")
        api_frame.pack(padx=10, pady=10, fill="x", expand=True, anchor="n") # Changed grid to pack
        
        api_inner_frame = ttk.Frame(api_frame) # Use an inner frame for grid layout
        api_inner_frame.pack(padx=5, pady=5, fill="x")

        ttk.Label(api_inner_frame, text="依赖状态:").grid(row=0, column=0, sticky="w", padx=5, pady=2) 
        ttk.Label(api_inner_frame, textvariable=self.api_dependencies_status).grid(row=0, column=1, columnspan=2, sticky="w", padx=5, pady=2)
        
        ttk.Label(api_inner_frame, text="API Key:").grid(row=1, column=0, sticky="w", padx=5, pady=2)
        api_entry = ttk.Entry(api_inner_frame, textvariable=self.deepseek_api_key, width=50, show="*")
        api_entry.grid(row=1, column=1, sticky="ew", padx=5, pady=2)
        
        def toggle_api_visibility():
            api_entry['show'] = '' if api_entry['show'] == '*' else '*'
            toggle_btn['text'] = '隐藏' if api_entry['show'] == '' else '显示'
        toggle_btn = ttk.Button(api_inner_frame, text="显示", width=5, command=toggle_api_visibility)
        toggle_btn.grid(row=1, column=2, padx=5, pady=2)

        ttk.Label(api_inner_frame, text="选择模型:").grid(row=2, column=0, sticky="w", padx=5, pady=2)
        model_selector = ttk.Combobox(api_inner_frame, textvariable=self.deepseek_model, values=["deepseek-chat", "deepseek-coder"], width=48, state="readonly")
        model_selector.grid(row=2, column=1, sticky="ew", padx=5, pady=2)

        save_api_btn = ttk.Button(api_inner_frame, text="保存API设置", command=self.save_api_settings)
        save_api_btn.grid(row=1, column=3, rowspan=2, padx=(10,5), pady=5, sticky="ns")
        api_inner_frame.columnconfigure(1, weight=1) # Make API key and model selector expand


        self.setup_ai_tab(ai_frame)
        self.setup_content_tab(content_frame)
        
        self.create_font_settings(title_settings_frame, 0, "文档标题", self.title_font, self.title_size, self.title_bold, self.title_color_val, "title_color")
        self.create_font_settings(title_settings_frame, 1, "一级标题", self.h1_font, self.h1_size, self.h1_bold, self.h1_color_val, "h1_color")
        self.create_font_settings(title_settings_frame, 2, "二级标题", self.h2_font, self.h2_size, self.h2_bold, self.h2_color_val, "h2_color")
        self.create_font_settings(title_settings_frame, 3, "三级标题", self.h3_font, self.h3_size, self.h3_bold, self.h3_color_val, "h3_color")
        
        self.create_font_settings(styles_frame, 0, "正文", self.normal_font, self.normal_size, self.normal_bold, self.normal_color_val, "normal_color")
        
        indent_frame = ttk.Frame(styles_frame)
        indent_frame.grid(row=1, column=0, columnspan=3, sticky="ew", padx=10, pady=10)
        ttk.Label(indent_frame, text="首行缩进字符数:").pack(side=tk.LEFT, padx=5)
        self.indent_entry = ttk.Entry(indent_frame, textvariable=self.indent_chars_str, width=5) 
        self.indent_entry.pack(side=tk.LEFT, padx=5)
        
        self.log_text_widget = tk.Text(log_frame, height=10, width=80, wrap=tk.WORD) 
        self.log_text_widget.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        self.log_text_widget.config(state=tk.DISABLED)
        scrollbar = ttk.Scrollbar(log_frame, command=self.log_text_widget.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.log_text_widget.config(yscrollcommand=scrollbar.set)
        
        self.progress_bar = ttk.Progressbar(self.root, orient="horizontal", length=980, mode="determinate")
        self.progress_bar.pack(pady=10)
        
        button_frame_main = ttk.Frame(self.root) 
        button_frame_main.pack(fill=tk.X, padx=10, pady=10)
        ttk.Button(button_frame_main, text="生成文档", command=self.generate_document).pack(side=tk.RIGHT, padx=5)
        
        self.init_default_sections()
        self.load_default_ui_settings() # Call again after widgets created to update color previews if they weren't ready
        self.toggle_logo_options_state() # Ensure logo options reflect loaded state

    def setup_header_footer_tab(self, parent_frame):
        logo_settings_frame = ttk.LabelFrame(parent_frame, text="页眉 Logo 设置")
        logo_settings_frame.pack(padx=10, pady=10, fill="x", anchor="n")

        add_logo_check = ttk.Checkbutton(logo_settings_frame, text="在页眉添加 Logo", variable=self.add_logo, command=self.toggle_logo_options_state)
        add_logo_check.grid(row=0, column=0, columnspan=3, sticky="w", padx=5, pady=5)

        self.logo_options_frame = ttk.Frame(logo_settings_frame)
        self.logo_options_frame.grid(row=1, column=0, columnspan=3, padx=5, pady=5, sticky="ew")
        logo_settings_frame.columnconfigure(1, weight=1)


        ttk.Label(self.logo_options_frame, text="Logo 图片路径:").grid(row=0, column=0, sticky="w", padx=5, pady=3)
        self.logo_path_entry = ttk.Entry(self.logo_options_frame, textvariable=self.logo_path, width=50, state="readonly")
        self.logo_path_entry.grid(row=0, column=1, sticky="ew", padx=5, pady=3)
        self.select_logo_button = ttk.Button(self.logo_options_frame, text="选择图片", command=self.select_logo_image_file)
        self.select_logo_button.grid(row=0, column=2, padx=5, pady=3)
        self.logo_options_frame.columnconfigure(1, weight=1)


        ttk.Label(self.logo_options_frame, text="Logo 位置:").grid(row=1, column=0, sticky="w", padx=5, pady=3)
        self.logo_pos_frame = ttk.Frame(self.logo_options_frame)
        self.logo_pos_frame.grid(row=1, column=1, columnspan=2, sticky="w", padx=5, pady=3)
        
        self.rb_logo_left = ttk.Radiobutton(self.logo_pos_frame, text="左", variable=self.logo_position, value="left")
        self.rb_logo_left.pack(side="left", padx=5)
        self.rb_logo_center = ttk.Radiobutton(self.logo_pos_frame, text="中", variable=self.logo_position, value="center")
        self.rb_logo_center.pack(side="left", padx=5)
        self.rb_logo_right = ttk.Radiobutton(self.logo_pos_frame, text="右", variable=self.logo_position, value="right")
        self.rb_logo_right.pack(side="left", padx=5)

        ttk.Label(self.logo_options_frame, text="Logo 宽度 (cm):").grid(row=2, column=0, sticky="w", padx=5, pady=3)
        self.logo_width_entry = ttk.Entry(self.logo_options_frame, textvariable=self.logo_width_cm, width=10)
        self.logo_width_entry.grid(row=2, column=1, sticky="w", padx=5, pady=3)
        
        # self.toggle_logo_options_state() # Called from create_widgets after all vars are set by load_default_ui_settings

    def toggle_logo_options_state(self):
        """Enables or disables logo option widgets based on the add_logo checkbox."""
        current_state = tk.NORMAL if self.add_logo.get() else tk.DISABLED
        
        # Check if widgets exist before configuring (important if called before full UI setup)
        if hasattr(self, 'logo_path_entry'):
            self.logo_path_entry.configure(state="readonly" if self.add_logo.get() else tk.DISABLED)
        if hasattr(self, 'select_logo_button'):
            self.select_logo_button.configure(state=current_state)
        if hasattr(self, 'rb_logo_left'): # Check one, assume others exist too
            self.rb_logo_left.configure(state=current_state)
            self.rb_logo_center.configure(state=current_state)
            self.rb_logo_right.configure(state=current_state)
        if hasattr(self, 'logo_width_entry'):
            self.logo_width_entry.configure(state=current_state)


    def select_logo_image_file(self):
        """Opens a file dialog to select a logo image."""
        if not self.add_logo.get():
            return
        filetypes = (
            ("Image files", "*.png *.jpg *.jpeg *.gif *.bmp *.tiff"),
            ("All files", "*.*")
        )
        filepath = filedialog.askopenfilename(title="选择 Logo 图片", filetypes=filetypes)
        if filepath:
            self.logo_path.set(filepath)
            self.log(f"已选择 Logo 图片: {filepath}")

    def init_default_sections(self):
        if not self.document_sections: # Only add if empty, e.g. on first run
            self.add_section(level=1, title="第一章 背景介绍", content="这里是背景介绍内容。") 
    
    def setup_content_tab(self, parent):
        paned_window = ttk.PanedWindow(parent, orient=tk.HORIZONTAL)
        paned_window.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        left_frame = ttk.Frame(paned_window); paned_window.add(left_frame, weight=1)
        right_frame = ttk.Frame(paned_window); paned_window.add(right_frame, weight=2)
        
        ttk.Label(left_frame, text="文档结构").pack(pady=5)
        self.tree = ttk.Treeview(left_frame, selectmode='browse')
        self.tree.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        self.tree["columns"] = ("title", "level")
        self.tree.column("#0", width=50, minwidth=50, stretch=tk.NO); self.tree.column("title", width=150, minwidth=100, stretch=tk.YES); self.tree.column("level", width=50, minwidth=50, stretch=tk.NO)
        self.tree.heading("#0", text="序号"); self.tree.heading("title", text="标题"); self.tree.heading("level", text="级别") 
        tree_scroll = ttk.Scrollbar(left_frame, orient="vertical", command=self.tree.yview); tree_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.tree.configure(yscrollcommand=tree_scroll.set)
        self.tree.bind('<<TreeviewSelect>>', self.on_tree_select)
        
        button_frame_tree = ttk.Frame(left_frame); button_frame_tree.pack(fill=tk.X, padx=5, pady=5)
        ttk.Button(button_frame_tree, text="添加章节", command=self.add_section_dialog).pack(side=tk.LEFT, padx=2)
        ttk.Button(button_frame_tree, text="编辑章节", command=self.edit_section_dialog).pack(side=tk.LEFT, padx=2)
        ttk.Button(button_frame_tree, text="删除章节", command=self.delete_section).pack(side=tk.LEFT, padx=2)
        ttk.Button(button_frame_tree, text="上移", command=lambda: self.move_section(-1)).pack(side=tk.LEFT, padx=2)
        ttk.Button(button_frame_tree, text="下移", command=lambda: self.move_section(1)).pack(side=tk.LEFT, padx=2)
        
        ttk.Label(right_frame, text="章节标题:").pack(anchor=tk.W, padx=5, pady=5)
        self.section_title_var = tk.StringVar()
        self.section_title_entry = ttk.Entry(right_frame, textvariable=self.section_title_var, width=40)
        self.section_title_entry.pack(fill=tk.X, padx=5, pady=2)
        
        level_frame = ttk.Frame(right_frame); level_frame.pack(fill=tk.X, padx=5, pady=5)
        ttk.Label(level_frame, text="章节级别:").pack(side=tk.LEFT)
        self.section_level_var = tk.IntVar(value=1)
        ttk.Radiobutton(level_frame, text="一级", variable=self.section_level_var, value=1).pack(side=tk.LEFT, padx=5)
        ttk.Radiobutton(level_frame, text="二级", variable=self.section_level_var, value=2).pack(side=tk.LEFT, padx=5)
        ttk.Radiobutton(level_frame, text="三级", variable=self.section_level_var, value=3).pack(side=tk.LEFT, padx=5)
        
        ttk.Label(right_frame, text="章节内容:").pack(anchor=tk.W, padx=5, pady=5)
        self.section_content_text = tk.Text(right_frame, wrap=tk.WORD, height=15)
        self.section_content_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        content_scroll = ttk.Scrollbar(right_frame, orient="vertical", command=self.section_content_text.yview); content_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.section_content_text.configure(yscrollcommand=content_scroll.set)
        
        bottom_frame_content = ttk.Frame(right_frame); bottom_frame_content.pack(fill=tk.X, padx=5, pady=5)
        ttk.Button(bottom_frame_content, text="保存章节", command=self.save_current_section).pack(side=tk.RIGHT, padx=5)
        self.current_section_id = None
    
    def add_section_dialog(self):
        dialog = tk.Toplevel(self.root); dialog.title("添加章节"); dialog.geometry("400x150"); dialog.transient(self.root); dialog.grab_set()
        ttk.Label(dialog, text="章节标题:").grid(row=0, column=0, sticky=tk.W, padx=10, pady=10)
        title_var = tk.StringVar(); ttk.Entry(dialog, textvariable=title_var, width=30).grid(row=0, column=1, padx=10, pady=10)
        ttk.Label(dialog, text="章节级别:").grid(row=1, column=0, sticky=tk.W, padx=10, pady=10)
        level_var = tk.IntVar(value=1)
        ttk.Radiobutton(dialog, text="一级", variable=level_var, value=1).grid(row=1, column=1, sticky=tk.W)
        ttk.Radiobutton(dialog, text="二级", variable=level_var, value=2).grid(row=1, column=1)
        ttk.Radiobutton(dialog, text="三级", variable=level_var, value=3).grid(row=1, column=1, sticky=tk.E)
        def on_confirm():
            title = title_var.get().strip(); level = level_var.get()
            if title: self.add_section(level, title); dialog.destroy()
            else: messagebox.showwarning("警告", "章节标题不能为空", parent=dialog)
        ttk.Button(dialog, text="确定", command=on_confirm).grid(row=2, column=1, sticky=tk.E, padx=10, pady=10)
        ttk.Button(dialog, text="取消", command=dialog.destroy).grid(row=2, column=0, sticky=tk.W, padx=10, pady=10)
        dialog.wait_window()
    
    def edit_section_dialog(self):
        selected = self.tree.selection()
        if not selected: messagebox.showinfo("提示", "请先选择要编辑的章节"); return
        section_id = selected[0]; section = self.find_section_by_id(section_id)
        if not section: return
        dialog = tk.Toplevel(self.root); dialog.title("编辑章节"); dialog.geometry("400x150"); dialog.transient(self.root); dialog.grab_set()
        ttk.Label(dialog, text="章节标题:").grid(row=0, column=0, sticky=tk.W, padx=10, pady=10)
        title_var = tk.StringVar(value=section['title']); ttk.Entry(dialog, textvariable=title_var, width=30).grid(row=0, column=1, padx=10, pady=10)
        ttk.Label(dialog, text="章节级别:").grid(row=1, column=0, sticky=tk.W, padx=10, pady=10)
        level_var = tk.IntVar(value=section['level'])
        ttk.Radiobutton(dialog, text="一级", variable=level_var, value=1).grid(row=1, column=1, sticky=tk.W)
        ttk.Radiobutton(dialog, text="二级", variable=level_var, value=2).grid(row=1, column=1)
        ttk.Radiobutton(dialog, text="三级", variable=level_var, value=3).grid(row=1, column=1, sticky=tk.E)
        def on_confirm():
            title = title_var.get().strip(); level = level_var.get()
            if title: section['title'] = title; section['level'] = level; self.update_tree(); dialog.destroy()
            else: messagebox.showwarning("警告", "章节标题不能为空", parent=dialog)
        ttk.Button(dialog, text="确定", command=on_confirm).grid(row=2, column=1, sticky=tk.E, padx=10, pady=10)
        ttk.Button(dialog, text="取消", command=dialog.destroy).grid(row=2, column=0, sticky=tk.W, padx=10, pady=10)
        dialog.wait_window()
    
    def add_section(self, level, title, content=""):
        section_id = f"section_{time.time_ns()}_{title.replace(' ','_')}"
        section = {'id': section_id, 'level': level, 'title': title, 'content': content}
        self.document_sections.append(section); self.update_tree(); return section['id']
    
    def update_tree(self):
        current_selection = self.tree.selection()
        current_focus = self.tree.focus()

        for item in self.tree.get_children(): self.tree.delete(item)
        for i, section in enumerate(self.document_sections):
            self.tree.insert("", tk.END, section['id'], text=str(i+1), values=(section['title'], f"级别{section['level']}"))
        
        if current_selection and self.tree.exists(current_selection[0]):
            self.tree.selection_set(current_selection[0])
        if current_focus and self.tree.exists(current_focus):
             self.tree.focus(current_focus)


    def find_section_by_id(self, section_id):
        for section in self.document_sections:
            if section['id'] == section_id: return section
        return None
    
    def setup_ai_tab(self, parent):
        ttk.Label(parent, text="粘贴您的文本内容，AI 将自动识别标题结构：").pack(anchor=tk.W, padx=10, pady=5)
        self.ai_input_text = scrolledtext.ScrolledText(parent, wrap=tk.WORD, height=15)
        self.ai_input_text.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        button_frame_ai = ttk.Frame(parent); button_frame_ai.pack(fill=tk.X, padx=10, pady=5)
        self.ai_status_var = tk.StringVar(value="准备就绪")
        status_label = ttk.Label(button_frame_ai, textvariable=self.ai_status_var)
        status_label.pack(side=tk.LEFT, padx=5)
        self.analyze_button = ttk.Button(button_frame_ai, text="识别标题并导入", command=self.analyze_with_deepseek)
        self.analyze_button.pack(side=tk.RIGHT, padx=5)

        tip_frame = ttk.LabelFrame(parent, text="使用说明"); tip_frame.pack(fill=tk.X, padx=10, pady=10)
        tips = ("1. 粘贴您的文本内容到上方文本框中\n2. 点击「识别标题并导入」按钮\n3. AI 将分析文本，识别各级标题\n"
                "4. 识别出的标题结构将自动导入到「文档内容」选项卡中\n5. 您可以在「文档内容」选项卡进一步编辑调整\n\n"
                "提示：标题识别最适合结构化文档，如学术论文、报告等带有明确章节标题的文本\n\n"
                "注意：首次使用时需要联网安装必要的依赖包。\n请确保您已经安装了 Python 并设置了正确的 DeepSeek API Key。")
        ttk.Label(tip_frame, text=tips, justify=tk.LEFT).pack(padx=10, pady=10)
    
    def analyze_with_deepseek(self):
        api_key = self.deepseek_api_key.get()
        if not api_key: messagebox.showerror("错误", "请先设置 DeepSeek API Key"); return
        text = self.ai_input_text.get(1.0, tk.END).strip()
        if not text: messagebox.showerror("错误", "请输入要分析的文本"); return
        
        if self.api_dependencies_status.get().startswith("缺失依赖"):
            messagebox.showwarning("依赖缺失", f"AI 功能所需依赖包缺失或检查失败: {self.api_dependencies_status.get().split(': ')[1]}\n请先确保依赖已正确安装。")
            return
        
        self.ai_status_var.set("正在分析中...") 
        if self.analyze_button: self.analyze_button.config(state=tk.DISABLED) 
        self.root.update_idletasks() 
        threading.Thread(target=self.run_deepseek_analysis, args=(text,), daemon=True).start()
    
    def run_deepseek_analysis(self, text):
        self.log("run_deepseek_analysis: 线程开始")
        api_key = self.deepseek_api_key.get()
        selected_model = self.deepseek_model.get() 
        self.log(f"run_deepseek_analysis: 使用模型: {selected_model}")

        api_url = "https://api.deepseek.com/chat/completions"
        headers = {"Content-Type": "application/json", "Authorization": f"Bearer {api_key}"}
        system_message_content = "你是一个专业的文本分析助手，负责识别文本中的标题结构，并严格按照用户指定的JSON格式返回结果。"
        user_prompt_content = f"""
请分析以下文本，识别其中的标题结构。将结果按如下JSON格式返回：
```json
[
    {{"level": 1, "title": "一级标题1", "content": "一级标题1下的正文内容"}},
    {{"level": 2, "title": "二级标题1.1", "content": "二级标题1.1下的正文内容"}},
    {{"level": 1, "title": "一级标题2", "content": "一级标题2下的正文内容"}},
    ...
]
```
规则：
1. level 表示标题级别，1为一级标题，2为二级标题，3为三级标题。
2. 识别标题时考虑格式特征，如数字编号（例如 1. 第一个, 1.1 小节, (一) 部分, A. 点）, 字体大小, 缩进等。
3. content 是标题下的正文内容，直到下一个同级或更高级别的标题出现之前的所有文本。如果标题下直接是子标题，则其 content 可以为空字符串。
4. 仅返回JSON格式，不要有其他解释文字或Markdown标记之外的内容。
5. 确保JSON是有效的，所有字符串都用双引号括起来。
以下是要分析的文本：
{text}
"""
        data = {"model": selected_model, "messages": [{"role": "system", "content": system_message_content}, {"role": "user", "content": user_prompt_content}]}
        
        max_retries = 3; backoff_factor = 0.5; request_timeout = 120 
        final_error_message = "AI分析失败，请检查网络连接和API Key。"

        for attempt in range(max_retries):
            self.log(f"run_deepseek_analysis: 尝试 {attempt + 1}/{max_retries} - 开始请求模型 {selected_model}")
            try:
                response = self.api_session.post(api_url, headers=headers, json=data, timeout=request_timeout)
                self.log(f"run_deepseek_analysis: 尝试 {attempt + 1} - 收到响应，状态码: {response.status_code}")

                if response.status_code == 200:
                    result = response.json()
                    if not result.get('choices') or not result['choices'][0].get('message') or not result['choices'][0]['message'].get('content'):
                        self.log(f"run_deepseek_analysis: 尝试 {attempt + 1} - API响应结构错误: {result}")
                        final_error_message = "AI未能生成有效响应内容或响应结构错误。"
                        if attempt == max_retries - 1: self.root.after(0, lambda msg=final_error_message: self.handle_ai_error(msg))
                        raise requests.exceptions.RequestException("API响应结构错误")

                    text_response = result['choices'][0]['message']['content']
                    if not text_response.strip():
                        self.log(f"run_deepseek_analysis: 尝试 {attempt + 1} - API返回空content")
                        final_error_message = "AI未能生成响应或响应为空。"
                        if attempt == max_retries - 1: self.root.after(0, lambda msg=final_error_message: self.handle_ai_error(msg))
                        raise requests.exceptions.RequestException("API返回空内容")

                    self.log(f"run_deepseek_analysis: 尝试 {attempt + 1} - 收到有效响应内容")
                    json_match = re.search(r'```json\s*([\s\S]*?)\s*```', text_response, re.DOTALL)
                    json_str = json_match.group(1).strip() if json_match else text_response.strip()
                    
                    try:
                        sections = json.loads(json_str)
                        self.log(f"run_deepseek_analysis: 尝试 {attempt + 1} - JSON解析成功，识别到 {len(sections)} 个章节。")
                        self.root.after(0, self.import_ai_sections, sections) 
                        return 
                    except json.JSONDecodeError as json_e:
                        self.log(f"run_deepseek_analysis: 尝试 {attempt + 1} - JSON解析错误: {json_e}. 内容: {json_str[:300]}...")
                        final_error_message = f"AI返回的JSON格式无效: {json_e}"
                        if attempt == max_retries - 1: self.root.after(0, lambda msg=final_error_message: self.handle_ai_error(msg))
                        raise requests.exceptions.RequestException("JSON解析错误")
                
                elif response.status_code >= 500: 
                    self.log(f"run_deepseek_analysis: 尝试 {attempt + 1} - API服务器错误 {response.status_code}: {response.text[:200]}")
                    final_error_message = f"AI请求失败(服务器错误)，错误码：{response.status_code}"
                    if attempt == max_retries - 1: self.root.after(0, lambda r=response, msg=final_error_message: self.handle_ai_error(f"{msg}\n详情: {r.text[:200]}..."))
                
                else: 
                    self.log(f"run_deepseek_analysis: 尝试 {attempt + 1} - API客户端错误 {response.status_code}: {response.text[:200]}")
                    error_detail = ""
                    try:
                        error_json = response.json()
                        error_detail = error_json.get("error", {}).get("message", response.text[:200])
                    except json.JSONDecodeError:
                        error_detail = response.text[:200]
                    final_error_message = f"AI请求失败(客户端错误)，错误码：{response.status_code}"
                    self.root.after(0, lambda msg=final_error_message, detail=error_detail: self.handle_ai_error(f"{msg}\n详情: {detail}"))
                    return 

            except requests.exceptions.Timeout as timeout_e:
                self.log(f"run_deepseek_analysis: 尝试 {attempt + 1} - 请求超时: {str(timeout_e)}")
                final_error_message = f"网络请求超时: {str(timeout_e)}"
                if attempt == max_retries - 1: self.root.after(0, lambda msg=final_error_message: self.handle_ai_error(msg))
            except requests.exceptions.RequestException as req_e: 
                self.log(f"run_deepseek_analysis: 尝试 {attempt + 1} - 网络或请求错误: {str(req_e)}")
                final_error_message = f"网络或请求错误: {str(req_e)}"
                if attempt == max_retries - 1: self.root.after(0, lambda msg=final_error_message: self.handle_ai_error(msg))
            
            if attempt < max_retries - 1:
                wait_time = backoff_factor * (2 ** attempt)
                self.log(f"run_deepseek_analysis: 等待 {wait_time:.2f} 秒后重试...")
                time.sleep(wait_time)
            elif attempt == max_retries -1: 
                self.log(f"run_deepseek_analysis: 所有 {max_retries} 次尝试均失败。最终错误: {final_error_message}")
        
        self.log("run_deepseek_analysis: 线程结束")
        if self.analyze_button and self.analyze_button.cget('state') == tk.DISABLED:
             self.root.after(0, lambda: self.analyze_button.config(state=tk.NORMAL))


    def import_ai_sections(self, sections_data):
        try:
            if not isinstance(sections_data, list): 
                self.handle_ai_error(f"AI返回的数据格式不正确，期望列表但得到 {type(sections_data)}。"); return
            
            if self.document_sections:
                if messagebox.askyesno("确认", "是否清空现有文档内容，并导入AI识别的章节？"):
                    self.document_sections = []; self.current_section_id = None
                    self.section_title_var.set(""); self.section_content_text.delete(1.0, tk.END)
                else:
                    self.ai_status_var.set("导入已取消")
                    if self.analyze_button: self.analyze_button.config(state=tk.NORMAL)
                    return

            imported_count = 0
            for section_item in sections_data:
                if not isinstance(section_item, dict): 
                    self.log(f"跳过无效的章节项目 (非字典): {section_item}"); continue
                title = section_item.get('title', f'未命名标题 {imported_count+1}')
                level = section_item.get('level')
                content = section_item.get('content', '')

                if not title or level is None:
                    self.log(f"跳过无效的章节项目 (缺少标题或级别): {section_item}"); continue
                try:
                    level = int(level)
                    if not (1 <= level <= 3):
                        self.log(f"跳过无效的章节项目 (级别超出范围1-3): {section_item}"); continue
                except ValueError:
                    self.log(f"跳过无效的章节项目 (级别非整数): {section_item}"); continue
                
                self.add_section(level=level, title=title, content=content)
                imported_count +=1
            
            self.update_tree()
            
            try: # Find the "文档内容" tab by its text attribute
                target_tab_text = "文档内容"
                target_tab_index = -1
                for i, tab_id in enumerate(self.notebook.tabs()):
                    if self.notebook.tab(tab_id, "text") == target_tab_text:
                        target_tab_index = i
                        break
                if target_tab_index != -1:
                    self.notebook.select(target_tab_index)
                else:
                    self.log(f"无法找到名为 '{target_tab_text}' 的标签页。")
            except Exception as e:
                self.log(f"切换到文档内容标签页时出错: {e}")


            self.ai_status_var.set(f"成功导入 {imported_count} 个章节")
            messagebox.showinfo("成功", f"已成功识别并导入 {imported_count} 个章节")
        except Exception as e: 
            self.handle_ai_error(f"导入章节时出错: {str(e)}\n原始数据: {str(sections_data)[:200]}...")
        finally: 
            if self.analyze_button: self.analyze_button.config(state=tk.NORMAL)
    
    def handle_ai_error(self, error_msg):
        self.ai_status_var.set("发生错误")
        self.log(f"AI错误: {error_msg}") 
        messagebox.showerror("AI 分析错误", error_msg)
        if self.analyze_button: self.analyze_button.config(state=tk.NORMAL) 
            
    def on_tree_select(self, event):
        selected_items = self.tree.selection()
        if not selected_items: return
        section_id = selected_items[0]; section = self.find_section_by_id(section_id)
        if section:
            if self.current_section_id and self.current_section_id != section_id: self.save_current_section() 
            self.current_section_id = section_id
            self.section_title_var.set(section['title']); self.section_level_var.set(section['level'])
            self.section_content_text.delete(1.0, tk.END); self.section_content_text.insert(tk.END, section['content'])
    
    def save_current_section(self):
        if not self.current_section_id: return
        section = self.find_section_by_id(self.current_section_id)
        if section:
            new_title, new_level, new_content = self.section_title_var.get().strip(), self.section_level_var.get(), self.section_content_text.get(1.0, tk.END).strip()
            if not new_title:
                messagebox.showwarning("警告", "章节标题不能为空。更改未保存。")
                self.section_title_var.set(section['title']) 
                return

            if (section['title'] != new_title or section['level'] != new_level or section['content'] != new_content):
                section.update({'title': new_title, 'level': new_level, 'content': new_content})
                self.update_tree(); self.log(f"章节 '{new_title}' 已保存。")
    
    def delete_section(self):
        selected = self.tree.selection()
        if not selected: messagebox.showinfo("提示", "请先选择要删除的章节"); return
        section_id = selected[0]
        section_to_delete = self.find_section_by_id(section_id)
        if messagebox.askyesno("确认", f"确定要删除章节 '{section_to_delete['title'] if section_to_delete else ''}' 吗？"):
            self.document_sections = [s for s in self.document_sections if s['id'] != section_id]
            self.update_tree()
            if self.current_section_id == section_id:
                self.current_section_id = None; self.section_title_var.set(""); self.section_level_var.set(1); self.section_content_text.delete(1.0, tk.END)
    
    def move_section(self, direction):
        selected = self.tree.selection()
        if not selected: messagebox.showinfo("提示", "请先选择要移动的章节"); return
        section_id = selected[0]
        for i, section_item in enumerate(self.document_sections):
            if section_item['id'] == section_id:
                new_pos = i + direction
                if 0 <= new_pos < len(self.document_sections):
                    self.document_sections[i], self.document_sections[new_pos] = self.document_sections[new_pos], self.document_sections[i]
                    self.update_tree(); self.tree.selection_set(section_id); self.tree.focus(section_id)
                break
    
    def create_font_settings(self, parent, r_idx, lbl_txt, fnt_var, sz_var, bld_var, clr_tk_var, clr_key):
        frame = ttk.LabelFrame(parent, text=lbl_txt); frame.grid(row=r_idx, column=0, columnspan=3, sticky="ew", padx=10, pady=10)
        frame.columnconfigure(1, weight=1) # Allow combobox to expand a bit if needed
        frame.columnconfigure(3, weight=1)

        ttk.Label(frame, text="字体:").grid(row=0, column=0, sticky="w", padx=5, pady=5)
        ttk.Combobox(frame, textvariable=fnt_var, values=["宋体", "黑体", "楷体", "仿宋", "微软雅黑", "Times New Roman", "Arial"], state="readonly", width=15).grid(row=0, column=1, sticky="ew", padx=5, pady=5)
        
        ttk.Label(frame, text="字号:").grid(row=0, column=2, sticky="w", padx=(10,5), pady=5)
        ttk.Combobox(frame, textvariable=sz_var, values=[str(s) for s in [8,9,10,10.5,11,12,14,15,16,18,20,21,22,24,26,28,36,42,48,72]], width=5, state="readonly").grid(row=0, column=3, sticky="w", padx=5, pady=5)
        
        ttk.Checkbutton(frame, text="加粗", variable=bld_var).grid(row=0, column=4, sticky="w", padx=5, pady=5)
        ttk.Button(frame, text="颜色", command=lambda k=clr_key: self.choose_color(k)).grid(row=0, column=5, sticky="w", padx=5, pady=5)
        
        # Initialize color preview with the current color from the variable
        initial_color = clr_tk_var.get() if clr_tk_var.get() else "#FFFFFF" # Default to white if empty
        preview = tk.Canvas(frame, width=20, height=20, bg=initial_color, highlightthickness=1, highlightbackground="black")
        preview.grid(row=0, column=6, sticky="w", padx=(5,10), pady=5)
        self.color_previews[clr_key] = preview
    
    def choose_color(self, color_key_name):
        color_tk_var = getattr(self, f"{color_key_name}_val", None)
        if not color_tk_var: self.log(f"错误：未找到颜色变量 {color_key_name}_val"); return
        
        initial_c = color_tk_var.get()
        chosen_color = colorchooser.askcolor(initialcolor=initial_c if initial_c else None, title="选择颜色")
        
        if chosen_color and chosen_color[1]: # chosen_color[1] is the hex string
            color_tk_var.set(chosen_color[1])
            if color_key_name in self.color_previews: 
                self.color_previews[color_key_name].config(bg=chosen_color[1])
            else: self.log(f"警告：未找到颜色预览 {color_key_name}")
    
    def log(self, message):
        if hasattr(self, 'log_text_widget') and self.log_text_widget:
            timestamp = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime())
            
            def _log(): # Use root.after to ensure thread safety for UI updates
                self.log_text_widget.config(state=tk.NORMAL)
                self.log_text_widget.insert(tk.END, f"[{timestamp}] {message}\n")
                self.log_text_widget.see(tk.END)
                self.log_text_widget.config(state=tk.DISABLED)
            
            if self.root.winfo_exists(): # Check if root window still exists
                 self.root.after(0, _log)
        else: print(f"LOG (widget not ready or root destroyed): {message}")
    
    def update_progress(self, value):
        if self.progress_bar and self.root.winfo_exists(): 
            self.root.after(0, lambda: self.progress_bar.config(value=value))

    def generate_document(self):
        self.save_current_section() 
        if not self.filename.get(): messagebox.showerror("错误", "请输入文件名称"); return
        if not self.document_sections: messagebox.showerror("错误", "文档内容为空，请至少添加一个章节"); return
        
        if self.add_logo.get():
            try:
                width = float(self.logo_width_cm.get())
                if width <= 0:
                    messagebox.showerror("错误", "Logo 宽度必须为正数。")
                    return
            except ValueError:
                messagebox.showerror("错误", "Logo 宽度必须是一个有效的数字。")
                return
            if not self.logo_path.get() or not os.path.exists(self.logo_path.get()):
                 messagebox.showerror("错误", "请选择一个有效的 Logo 图片路径。")
                 return

        # Prompt for save directory in the main thread before starting the generation thread
        save_dir = filedialog.askdirectory(title="选择保存文档的文件夹")
        if not save_dir: 
            self.log("用户取消了文件夹选择，操作终止。")
            messagebox.showinfo("取消", "文档生成已取消。")
            return
            
        doc_path = os.path.join(save_dir, f"{self.filename.get().strip()}.docx")
        
        # Disable generate button during generation
        # Assuming self.generate_button is the main generate button
        # You might need to find it or pass it if it's not a direct attribute
        # For now, let's assume there's a way to disable it.
        # Example: if hasattr(self, 'main_generate_button'): self.main_generate_button.config(state=tk.DISABLED)

        threading.Thread(target=self.generate_document_thread, args=(doc_path,), daemon=True).start()
    
    def apply_header_settings(self, document):
        if not self.add_logo.get() or not self.logo_path.get():
            self.log("未选择添加 Logo 或未指定 Logo 图片路径，跳过页眉 Logo 设置。")
            return

        logo_path_str = self.logo_path.get()
        if not os.path.exists(logo_path_str):
            self.log(f"Logo 图片路径无效: {logo_path_str}")
            return

        self.log(f"开始添加 Logo 到页眉: {logo_path_str}")
        position = self.logo_position.get()
        try:
            width_cm_val = float(self.logo_width_cm.get())
            if width_cm_val <=0:
                self.log(f"Logo 宽度无效 ({width_cm_val}cm)，将使用默认宽度 2.5cm。")
                width_cm_val = 2.5
        except ValueError:
            self.log(f"Logo 宽度值无效，将使用默认宽度 2.5cm。")
            width_cm_val = 2.5


        for section in document.sections:
            header = section.header
            
            while header.paragraphs:
                p_to_remove = header.paragraphs[0]
                header._element.remove(p_to_remove._element)

            logo_paragraph = header.add_paragraph()

            if position == "left":
                logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif position == "center":
                logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif position == "right":
                logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else: 
                logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                self.log(f"未知的 Logo 位置 '{position}'，默认为左对齐。")
            
            try:
                run = logo_paragraph.add_run()
                run.add_picture(logo_path_str, width=Cm(width_cm_val))
                self.log(f"Logo 已添加到页眉，位置: {position}, 宽度: {width_cm_val}cm")
            except FileNotFoundError:
                self.log(f"错误: Logo 文件未找到于 '{logo_path_str}'。Logo 未添加。")
            except Exception as e:
                self.log(f"添加 Logo 图片时出错: {e}")

    def generate_document_thread(self, doc_path): # doc_path is now passed as argument
        try:
            # Clear log widget from main thread using root.after
            if self.log_text_widget and self.root.winfo_exists():
                self.root.after(0, lambda: [
                    self.log_text_widget.config(state=tk.NORMAL),
                    self.log_text_widget.delete(1.0, tk.END),
                    self.log_text_widget.config(state=tk.DISABLED)
                ])

            self.log("开始文档生成过程..."); self.update_progress(5)
            self.log(f"文档将保存至: {doc_path}"); self.update_progress(10)
            
            doc = Document(); self.log("已创建新文档..."); self.update_progress(15)
            
            self.log("设置文档页面格式...")
            for sec in doc.sections: 
                sec.page_height,sec.page_width=Inches(11.69),Inches(8.27) 
                sec.left_margin,sec.right_margin,sec.top_margin,sec.bottom_margin=Inches(1.25),Inches(1.0),Inches(1.0),Inches(1.0) 
            self.update_progress(20)

            self.log("应用页眉设置...")
            self.apply_header_settings(doc)
            self.update_progress(25) 
            
            self.log("创建文档样式...")
            self.create_style(doc,'DocTitleStyle','文档标题',self.title_font.get(),self.title_size.get(),self.title_bold.get(),self.title_color_val.get()); self.update_progress(30)
            self.create_style(doc,'Heading1Style','一级标题',self.h1_font.get(),self.h1_size.get(),self.h1_bold.get(),self.h1_color_val.get(),level=1); self.update_progress(40)
            self.create_style(doc,'Heading2Style','二级标题',self.h2_font.get(),self.h2_size.get(),self.h2_bold.get(),self.h2_color_val.get(),level=2); self.update_progress(50)
            self.create_style(doc,'Heading3Style','三级标题',self.h3_font.get(),self.h3_size.get(),self.h3_bold.get(),self.h3_color_val.get(),level=3); self.update_progress(60)
            
            self.log("创建正文样式...")
            n_style=doc.styles['Normal']; n_style.font.name=self.normal_font.get(); n_style.font.size=Pt(self.normal_size.get()); n_style.font.bold=self.normal_bold.get()
            n_style.font.color.rgb=self.rgb_from_hex(self.normal_color_val.get()); 
            
            # Ensure East Asian font for Normal style as well
            rFonts_normal = n_style.element.rPr.get_or_add_rFonts()
            rFonts_normal.set(qn('w:eastAsia'),self.normal_font.get())
            
            try: 
                indent_chars_val = int(self.indent_entry.get()) # Get from the Entry widget
            except ValueError: 
                self.log("警告：首行缩进字符数无效，默认为0。"); indent_chars_val=0
            
            # Use a more standard indent calculation: font size * number of characters
            # For Chinese characters, this is a reasonable approximation.
            indent_val = Pt(indent_chars_val * self.normal_size.get()) if indent_chars_val > 0 else Pt(0)

            n_style.paragraph_format.first_line_indent = indent_val
            n_style.paragraph_format.line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE; self.update_progress(70)
            
            self.log("添加文档内容..."); self.log("添加文档标题...")
            title_p=doc.add_paragraph(self.document_title.get(),style='DocTitleStyle'); title_p.alignment=WD_ALIGN_PARAGRAPH.CENTER; self.update_progress(75)
            
            self.log("添加目录..."); self.add_toc(doc,self.toc_title.get()); doc.add_page_break(); self.update_progress(80)
            self.log("添加文档主体内容..."); self.add_user_document_content(doc); self.update_progress(90)
            
            self.log("保存文档..."); doc.save(doc_path); self.update_progress(100)
            self.log(f"文档已成功保存至 {doc_path}"); self.log("请在Word中打开文档，右键点击目录，选择'更新域'或按F9更新目录。")
            if self.root.winfo_exists():
                self.root.after(0,lambda: messagebox.showinfo("成功",f"文档已成功生成并保存至:\n{doc_path}\n\n请在Word中打开文档后，右键点击目录，选择'更新域'更新目录。"))
        except Exception as e:
            err_msg=f"生成文档时发生错误:\n{str(e)}"; self.log(f"错误: {str(e)}")
            import traceback
            self.log(f"Traceback: {traceback.format_exc()}")
            if self.root.winfo_exists():
                self.root.after(0,lambda em=err_msg: messagebox.showerror("错误",em)); 
            self.update_progress(0)
        finally:
            # Re-enable generate button
            # if hasattr(self, 'main_generate_button') and self.root.winfo_exists(): 
            #    self.root.after(0, lambda: self.main_generate_button.config(state=tk.NORMAL))
            pass # Placeholder for re-enabling button if needed

    def add_user_document_content(self, document):
        self.log("开始添加用户定义的文档内容...")
        for sec_item in self.document_sections:
            self.log(f"添加章节: {sec_item['title']} (级别 {sec_item['level']})"); style_name='Normal'
            if sec_item['level']==1: style_name='Heading1Style'
            elif sec_item['level']==2: style_name='Heading2Style'
            elif sec_item['level']==3: style_name='Heading3Style'
            
            try: 
                p = document.add_paragraph(sec_item['title'],style=style_name)
                if style_name.startswith('Heading'):
                    p.paragraph_format.first_line_indent = None 
            except KeyError: 
                self.log(f"警告：样式 '{style_name}' 未找到，使用 Normal 样式替代。"); 
                p = document.add_paragraph(sec_item['title'],style='Normal')

            if sec_item['content']:
                content_paragraphs = re.split(r'\n\s*\n', sec_item['content'].strip())
                for para_text in content_paragraphs:
                    if para_text.strip():
                        lines = para_text.splitlines()
                        if lines:
                            first_line_para = document.add_paragraph(style='Normal')
                            # Add the first line with potential indent from 'Normal' style
                            first_line_para.add_run(lines[0].strip())
                            
                            # For subsequent lines of the same original paragraph, add them without indent
                            # by adding them as runs to the same paragraph, or as new paragraphs
                            # where first_line_indent is explicitly removed or a non-indented style is used.
                            for line_text in lines[1:]:
                                if line_text.strip():
                                     # Simplest: add as a new paragraph with Normal style.
                                     # If Normal style has first_line_indent, this might not be desired for continuation lines.
                                     # A more complex approach would be to create a 'BodyTextContinuation' style without indent.
                                     # Or, add a soft line break (Shift+Enter) if supported, or just new paragraphs.
                                     document.add_paragraph(line_text.strip(), style='Normal') 
        self.log("所有用户定义的内容已添加完成")
    
    def create_style(self,document,style_id,style_name_ui,font_name,font_size_pt,is_bold,color_hex,level=None):
        try:
            self.log(f"创建样式: {style_name_ui}, 字体: {font_name}, 大小: {font_size_pt}pt, 加粗: {is_bold}, 颜色: {color_hex}")
            try: style=document.styles[style_id]
            except KeyError: style=document.styles.add_style(style_id,WD_STYLE_TYPE.PARAGRAPH)
            
            style.name = style_name_ui
            style.hidden = False
            style.quick_style = True 

            font=style.font
            font.name = font_name
            try:
                font.size = Pt(float(font_size_pt)) 
            except ValueError:
                self.log(f"字号 '{font_size_pt}' 无效，将使用默认值 12pt for style {style_id}")
                font.size = Pt(12)

            font.bold = is_bold
            font.color.rgb = self.rgb_from_hex(color_hex)
            
            rpr = font.element.get_or_add_rPr() # Ensure rPr exists
            rpr_fonts = rpr.get_or_add_rFonts() # Ensure rFonts exists

            rpr_fonts.set(qn('w:eastAsia'),font_name)
            rpr_fonts.set(qn('w:ascii'),font_name) 
            rpr_fonts.set(qn('w:hAnsi'),font_name)
            
            p_fmt=style.paragraph_format
            if level: 
                try: 
                    base_style_name=f'Heading {level}'
                    style.base_style = document.styles[base_style_name] if base_style_name in document.styles else document.styles['Normal']
                except KeyError: 
                    self.log(f"警告: 内置标题样式 'Heading {level}' 未找到，基于 Normal 创建。")
                    style.base_style=document.styles['Normal']
                
                style.next_paragraph_style=document.styles['Normal']
                pPr = style.element.get_or_add_pPr()
                
                # Remove existing outlineLvl if present before adding new one to avoid duplicates
                existing_outlineLvl = pPr.find(qn('w:outlineLvl'))
                if existing_outlineLvl is not None:
                    pPr.remove(existing_outlineLvl)

                outlineLvl = OxmlElement('w:outlineLvl')
                outlineLvl.set(qn('w:val'), str(level - 1)) 
                pPr.append(outlineLvl)

                p_fmt.space_before = Pt(12 if level == 1 else (8 if level == 2 else 6))
                p_fmt.space_after = Pt(6 if level == 1 else (4 if level == 2 else 2))
                p_fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE 
                p_fmt.alignment=WD_ALIGN_PARAGRAPH.LEFT
                p_fmt.keep_with_next = True 
                p_fmt.keep_together = True 
                p_fmt.first_line_indent = None 

            elif style_id == 'DocTitleStyle': 
                p_fmt.space_before=Pt(18)
                p_fmt.space_after=Pt(18)
                p_fmt.alignment=WD_ALIGN_PARAGRAPH.CENTER
                p_fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p_fmt.keep_with_next = True
            
            return style
        except Exception as e: 
            self.log(f"创建样式 '{style_name_ui}' 时出错: {str(e)}")
            import traceback
            self.log(f"Traceback: {traceback.format_exc()}")
            return None
    
    def rgb_from_hex(self,hex_color_str):
        hex_color_str=hex_color_str.lstrip('#')
        if len(hex_color_str)==6: 
            try:
                return RGBColor(int(hex_color_str[0:2],16),int(hex_color_str[2:4],16),int(hex_color_str[4:6],16))
            except ValueError:
                self.log(f"警告：无效的十六进制颜色值在 '{hex_color_str}' 中，使用黑色。")
                return RGBColor(0,0,0)
        else: self.log(f"警告：无效的十六进制颜色字符串 '{hex_color_str}'，使用黑色。"); return RGBColor(0,0,0)
    
    def add_toc(self,document,toc_main_title="目录"):
        try:
            self.log(f"添加目录标题: {toc_main_title}")
            p_title=document.add_paragraph(); 
            p_title.alignment=WD_ALIGN_PARAGRAPH.CENTER
            run_title=p_title.add_run(toc_main_title)
            run_title.font.name=self.h1_font.get() 
            try:
                toc_title_size = self.h1_size.get()
                run_title.font.size=Pt(toc_title_size if toc_title_size > 16 else 16)
            except tk.TclError: # Handle case where h1_size might not be a valid int yet
                 run_title.font.size=Pt(16)
                 self.log("警告: H1字号无效，目录标题字号设为16pt")

            run_title.font.bold=True # TOC title usually bold
            run_title.font.color.rgb = self.rgb_from_hex(self.h1_color_val.get()) 
            
            rpr_title = run_title.font.element.get_or_add_rPr()
            rpr_fonts_title = rpr_title.get_or_add_rFonts()
            rpr_fonts_title.set(qn('w:eastAsia'),self.h1_font.get())
            
            p_title.paragraph_format.space_before = Pt(12)
            p_title.paragraph_format.space_after=Pt(12)

            self.log("添加目录字段")
            paragraph = document.add_paragraph()
            run = paragraph.add_run()
            fldChar_begin = OxmlElement('w:fldChar')
            fldChar_begin.set(qn('w:fldCharType'), 'begin')
            
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = r' TOC \o "1-3" \h \z \u ' 
            
            fldChar_separate = OxmlElement('w:fldChar')
            fldChar_separate.set(qn('w:fldCharType'), 'separate')

            fldChar_end = OxmlElement('w:fldChar')
            fldChar_end.set(qn('w:fldCharType'), 'end')

            run._r.append(fldChar_begin)
            run._r.append(instrText)
            run._r.append(fldChar_separate)
            run._r.append(fldChar_end)
            
            self.log("目录添加完成"); return paragraph
        except Exception as e: 
            self.log(f"添加目录时出错: {str(e)}")
            import traceback
            self.log(f"Traceback: {traceback.format_exc()}")
            return None

if __name__ == "__main__":
    root = tk.Tk()
    app = DocxFormatter(root)
    root.mainloop()
