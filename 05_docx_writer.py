from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

class DocxWriter:
    def __init__(self, style_config, content_manager, log_callback, progress_callback):
        self.style_config = style_config
        self.content_manager = content_manager
        self.log = log_callback
        self.update_progress = progress_callback

    def _convert_hex_to_rgb(self, hex_color):
        """将十六进制颜色值转换为RGB对象"""
        hex_color = hex_color.lstrip('#')
        r = int(hex_color[0:2], 16)
        g = int(hex_color[2:4], 16)
        b = int(hex_color[4:6], 16)
        return RGBColor(r, g, b)

    def _create_or_get_style(self, document, style_id, style_name, font_name, font_size_pt, 
                               is_bold, color_hex, east_asia_font, level=None):
        try:
            self.log(f"创建或获取样式: {style_name}, 字体: {font_name}, 大小: {font_size_pt}pt")
            
            if style_id in document.styles:
                style = document.styles[style_id]
            else:
                style = document.styles.add_style(style_id, WD_STYLE_TYPE.PARAGRAPH)
            
            style.name = style_name
            style.hidden = False
            style.quick_style = True
            
            if level:
                # base_style_name = 'Heading ' + str(level) # docx library might not have this directly
                # if base_style_name in document.styles:
                # style.base_style = document.styles[base_style_name]
                # else:
                #     self.log(f"警告: 基础样式 '{base_style_name}' 未找到，将不设置基础样式.")
                style.next_paragraph_style = document.styles['Normal']
                p_pr = style._element.get_or_add_pPr()
                # Ensure numPr and ilvl elements are created correctly for TOC
                num_pr = p_pr.get_or_add_numPr()
                ilvl = num_pr.get_or_add_ilvl()
                ilvl.set(qn('w:val'), str(level - 1))
                num_id = num_pr.get_or_add_numId()
                num_id.set(qn('w:val'), '0') # Using a default numId, can be more specific if needed

            font = style.font
            font.name = font_name
            font.size = Pt(font_size_pt)
            font.bold = is_bold
            font.color.rgb = self._convert_hex_to_rgb(color_hex)
            font._element.rPr.rFonts.set(qn('w:eastAsia'), east_asia_font)
            
            paragraph_format = style.paragraph_format
            if style_id == 'Title':
                paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph_format.space_before = Pt(12) # Example, adjust as needed
                paragraph_format.space_after = Pt(12)  # Example, adjust as needed
            elif level: # Heading styles
                paragraph_format.space_before = Pt(12) 
                paragraph_format.space_after = Pt(4)
                paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            
            return style
        except Exception as e:
            self.log(f"创建样式 '{style_name}' 时出错: {str(e)}")
            return None

    def _create_styles_in_document(self, document):
        self.log("创建文档样式...")
        sc = self.style_config

        # 文档标题样式
        self._create_or_get_style(document, 'Title', '文档标题', 
                                  sc.title_font_var.get(), sc.title_size_var.get(), 
                                  sc.title_bold_var.get(), sc.title_color_hex, sc.title_font_var.get())
        self.update_progress(30)

        # 一级标题样式
        self._create_or_get_style(document, 'Heading1', '一级标题', 
                                  sc.h1_font_var.get(), sc.h1_size_var.get(), 
                                  sc.h1_bold_var.get(), sc.h1_color_hex, sc.h1_font_var.get(), level=1)
        self.update_progress(40)

        # 二级标题样式
        self._create_or_get_style(document, 'Heading2', '二级标题', 
                                  sc.h2_font_var.get(), sc.h2_size_var.get(), 
                                  sc.h2_bold_var.get(), sc.h2_color_hex, sc.h2_font_var.get(), level=2)
        self.update_progress(50)

        # 三级标题样式
        self._create_or_get_style(document, 'Heading3', '三级标题', 
                                  sc.h3_font_var.get(), sc.h3_size_var.get(), 
                                  sc.h3_bold_var.get(), sc.h3_color_hex, sc.h3_font_var.get(), level=3)
        self.update_progress(60)

        # 正文样式
        self.log("创建正文样式...")
        normal_style = document.styles['Normal']
        normal_style.font.name = sc.normal_font_var.get()
        normal_style.font.size = Pt(sc.normal_size_var.get())
        normal_style.font.bold = sc.normal_bold_var.get()
        normal_style.font.color.rgb = self._convert_hex_to_rgb(sc.normal_color_hex)
        normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), sc.normal_font_var.get())
        
        try:
            indent_val = int(sc.indent_chars_var.get())
            normal_style.paragraph_format.first_line_indent = Pt(indent_val * sc.normal_size_var.get()) 
        except ValueError:
            self.log(f"错误: 首行缩进字符数 '{sc.indent_chars_var.get()}' 不是有效数字。将不设置首行缩进。")
            normal_style.paragraph_format.first_line_indent = None
        normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        self.update_progress(70)

    def _add_table_of_contents(self, document, toc_title_str):
        try:
            self.log(f"添加目录标题: {toc_title_str}")
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(toc_title_str)
            # Apply H1 style attributes for TOC title as per original logic for font consistency
            run.font.name = self.style_config.h1_font_var.get()
            run.font.size = Pt(self.style_config.h1_size_var.get()) # Using H1 size for TOC title as in original
            run.font.bold = self.style_config.h1_bold_var.get()
            run._element.rPr.rFonts.set(qn('w:eastAsia'), self.style_config.h1_font_var.get())
            
            self.log("添加目录字段")
            p = document.add_paragraph()
            run = p.add_run()
            
            fldChar = OxmlElement('w:fldChar')
            fldChar.set(qn('w:fldCharType'), 'begin')
            
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = r' TOC \o "1-3" \h \z \u '

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'separate')
            
            fldChar3 = OxmlElement('w:fldChar')
            fldChar3.set(qn('w:fldCharType'), 'end')
            
            r_element = run._element
            r_element.append(fldChar)
            r_element.append(instrText)
            r_element.append(fldChar2)
            r_element.append(fldChar3)
            
            self.log("目录添加完成")
        except Exception as e:
            self.log(f"添加目录时出错: {str(e)}")

    def _add_user_defined_content(self, document):
        self.log("开始添加用户定义的文档内容...")
        all_sections = self.content_manager.get_all_sections()
        
        for section in all_sections:
            self.log(f"添加章节: {section['title']} (级别 {section['level']})")
            style_name = f"Heading{section['level']}" if section['level'] in [1,2,3] else 'Normal'
            
            try:
                document.add_paragraph(section['title'], style=style_name)
            except KeyError:
                self.log(f"警告: 样式 '{style_name}' 未找到，将使用Normal样式添加标题 '{section['title']}'")
                document.add_paragraph(section['title'], style='Normal')

            if section['content']:
                # Ensure content is a string and not None
                content_text = str(section['content']).strip()
                if content_text: # Only add paragraph if there's actual content
                    document.add_paragraph(content_text, style='Normal')
        
        self.log("所有用户定义的内容已添加完成")

    def generate_document(self, output_dir, filename_str, document_title_str, toc_title_str):
        try:
            self.log("开始文档生成过程...")
            self.update_progress(5)
            
            doc_path = os.path.join(output_dir, f"{filename_str}.docx")
            self.log(f"文档将保存至: {doc_path}")
            self.update_progress(10)
            
            document = Document()
            self.log("已创建新文档...")
            self.update_progress(15)
            
            self.log("设置文档页面格式...")
            doc_sections = document.sections
            for section_props in doc_sections:
                section_props.page_height = Inches(11.69)
                section_props.page_width = Inches(8.27)
                section_props.left_margin = Inches(1)
                section_props.right_margin = Inches(1)
                section_props.top_margin = Inches(1)
                section_props.bottom_margin = Inches(0.8)
            self.update_progress(20)
            
            self._create_styles_in_document(document)
            # Progress already updated within _create_styles_in_document
            
            self.log("添加文档标题...")
            if document_title_str:
                title_para = document.add_paragraph(document_title_str, style='Title')
                # Alignment already handled by style
            self.update_progress(75)
            
            self._add_table_of_contents(document, toc_title_str)
            document.add_page_break()
            self.update_progress(80)
            
            self._add_user_defined_content(document)
            self.update_progress(90)
            
            self.log("保存文档...")
            document.save(doc_path)
            self.update_progress(100)
            
            self.log(f"文档已成功保存至 {doc_path}")
            self.log("请在Word中打开文档，右键点击目录，选择'更新域'或按F9更新目录。")
            return doc_path, None # path, error
        except Exception as e:
            error_message = f"生成文档时发生错误: {str(e)}"
            self.log(f"错误: {error_message}")
            self.update_progress(0)
            return None, error_message 